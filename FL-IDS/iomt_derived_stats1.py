from docx import Document

# Create a new Word document for Device Metric Summary Statistics
doc = Document()
doc.add_heading("Device Metric Summary Statistics", level=1)

# Define table data
data = [
    ["Device", "Metric", "Mean", "SD", "95% CI", "Median", "Min", "Max"],
    ["WIP", "TRR", "89.69", "8.09", "[83.47, 95.90]", "86.92", "75.11", "100.00"],
    ["WIP", "PLR", "50.63", "33.15", "[25.15, 76.11]", "71.14", "0.00", "80.98"],
    ["WIP", "NDI", "49.37", "33.15", "[23.89, 74.85]", "28.86", "19.02", "100.00"],
    ["WIP", "JVI", "45.84", "44.26", "[11.82, 79.86]", "31.60", "0.41", "100.00"],
    ["WIP", "ASI", "57.44", "17.65", "[43.88, 71.01]", "51.48", "41.84", "100.00"],
    ["SHS", "TRR", "89.43", "8.29", "[83.06, 95.80]", "86.60", "74.49", "100.00"],
    ["SHS", "PLR", "56.55", "36.96", "[28.14, 84.96]", "73.54", "0.00", "92.92"],
    ["SHS", "NDI", "43.45", "36.96", "[15.04, 71.86]", "26.46", "7.08", "100.00"],
    ["SHS", "JVI", "41.64", "45.35", "[6.78, 76.50]", "19.44", "0.41", "100.00"],
    ["SHS", "ASI", "46.84", "27.49", "[25.71, 67.97]", "34.21", "24.22", "100.00"],
]

# Create table
table = doc.add_table(rows=1, cols=len(data[0]))
table.style = "Table Grid"

# Add header
hdr_cells = table.rows[0].cells
for i, header in enumerate(data[0]):
    hdr_cells[i].text = header

# Add rows
for row_data in data[1:]:
    row_cells = table.add_row().cells
    for i, item in enumerate(row_data):
        row_cells[i].text = item

# Save the document
output_path = "/home/ramamr36/ns-allinone-3.42/ns-3.42/Device_Metric_Summary_Statistics.docx"
doc.save(output_path)

output_path

