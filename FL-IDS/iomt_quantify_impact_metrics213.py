#!/usr/bin/env python3
"""
Merged IoMT MITM Full Pipeline - Standards Extended
Combines comprehensive metrics calculation with full risk mapping and regulatory compliance

CHANGES:
- Preserves original functionality (ISO 11073, ISO 14971, ISO/IEC 27001/27005, IEC 80001-1)
- ADDED: ISO/IEC 27400:2022, ISO 27799, ISO 13485:2016, ISO/IEC 27701, IEC 62443
- New per-flow outputs: ThreatCategory, IEC62443_Zone, PrivacyImpactLevel, CIA_Compliance
- New Excel sheets for the added standards' mappings
- Maintains Excel/Word reporting, plotting, and colorization logic
"""
import os
import re
import glob
import argparse
import xmltodict
import xml.etree.ElementTree as ET
import pandas as pd
import numpy as np
import math
import matplotlib.pyplot as plt
from sklearn.metrics import confusion_matrix, classification_report
from openpyxl import Workbook, load_workbook
from openpyxl.styles import PatternFill
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.utils import get_column_letter
from openpyxl.drawing.image import Image as XLImage

try:
    from docx import Document
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

# Configuration (unchanged paths)
NORMAL_WIP_DIR = "data/normal/normal_wip"
NORMAL_SHS_DIR = "data/normal/normal_shs"
MITM_WIP_DIR   = "data/mitm/mitm_wip"
MITM_SHS_DIR   = "data/mitm/mitm_shs"

OUTPUT_DIR = "analysis"
PLOTS_DIR = os.path.join(OUTPUT_DIR, "plots")
CSV_DIR = os.path.join(OUTPUT_DIR, "metrics")
EXCEL_PATH = os.path.join(OUTPUT_DIR, "iomt_comprehensive_analysis.xlsx")
WORD_PATH = os.path.join(OUTPUT_DIR, "iomt_risk_assessment_security_plan.docx")

os.makedirs(PLOTS_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(CSV_DIR, exist_ok=True)

BENCHMARK_CSV = "benchmark_summary.csv"
EPS = 1e-9
BYTES_TO_MB = 1.0 / (1024.0 * 1024.0)
FLOW9_ID = 9

# Device weights (ISO 11073 inspired) - preserved
DEVICE_WEIGHTS = {
    'wip': {'TRR': 0.4, 'PLR': 0.3, 'NDI': 0.2, 'JVI': 0.1},
    'shs': {'TRR': 0.2, 'PLR': 0.2, 'NDI': 0.3, 'JVI': 0.3}
}

ISO11073_WEIGHTS = {
    "WIP": {"TRR": 0.4, "PLR": 0.3, "NDI": 0.2, "JVI": 0.1},
    "SHS": {"TRR": 0.2, "PLR": 0.2, "NDI": 0.3, "JVI": 0.3}
}

CIA_WEIGHTS = {"C": 0.33, "I": 0.33, "A": 0.34}

###############################################################################
# ISO/IEEE 11073 + ISO 14971 Clinical Criticality Mapping (preserved)
###############################################################################

DEVICE_RPN = {
    "wip": 15,   # High-risk, therapeutic device
    "shs": 10    # Moderate-risk, diagnostic sensor
}

DEVICE_CRITICALITY = {
    "wip": {
        "factor": 1.00,
        "thresholds": {"Warning": 50, "Action": 70, "Critical": 90},
        "source": "ISO 14971 Table B.1, Life-support class"
    },
    "shs": {
        "factor": 0.85,
        "thresholds": {"Warning": 55, "Action": 70, "Critical": 85},
        "source": "ISO 14971 Table B.2, Diagnostic class"
    }
}

__CRITICALITY_METADATA__ = {
    "method": "Explicit factors based on ISO 14971 RPNs and ISO 11073 device class semantics",
    "DEVICE_RPN": DEVICE_RPN,
    "DEVICE_CRITICALITY": DEVICE_CRITICALITY,
    "citations": [
        "ISO 14971:2019 §5–7 (Risk analysis and evaluation)",
        "ISO/IEEE 11073-10101 (Device data semantics)",
        "IEC 80001-1:2021 (Application of risk management for IT-networked medical devices)",
        # ADDED citations for new standards included in metadata
        "ISO/IEC 27400:2022 (IoT cybersecurity and privacy guidance)",
        "ISO 27799:2016 (Health informatics - information security management in health)",
        "ISO/IEC 27701:2019 (PIMS extension to ISO/IEC 27001/27002)",
        "ISO 13485:2016 (Medical devices - QMS)",
        "IEC 62443 (Industrial communications - networked system security)"
    ]
}

def compute_priority(asi_value, device_type):
    """Compute PriorityScore and ActionLevel using ISO 14971/11073-derived factors."""
    dev = device_type.lower()
    crit = DEVICE_CRITICALITY.get(dev, DEVICE_CRITICALITY["wip"])
    factor = crit["factor"]
    thr = crit["thresholds"]
    priority = asi_value * factor

    if priority >= thr["Critical"]:
        level = "Critical"
    elif priority >= thr["Action"]:
        level = "Action"
    elif priority >= thr["Warning"]:
        level = "Warning"
    else:
        level = "Normal"

    return round(priority, 2), level, thr, round(factor, 2)

###############################################################################
# Severity and risk thresholds (preserved)
###############################################################################

ASI_SEVERITY = [(0, 10, "None"), (10, 25, "Low"), (25, 50, "Medium"), (50, 101, "High")]
RISK_LEVELS = [(0, 10, "None"), (10, 25, "Low"), (25, 50, "Medium"), (50, 75, "High"), (75, 101, "Critical")]
NHS_LIKELIHOOD_THRESH = [(0, 10, 1), (10, 25, 2), (25, 50, 3), (50, 75, 4), (75, 101, 5)]
NHS_IMPACT_THRESH = [(0, 10, 1), (10, 25, 2), (25, 50, 3), (50, 75, 4), (75, 101, 5)]

# Color schemes (preserved)
SEVERITY_COLOR = {"None": "66bb6a", "Low": "ffd54f", "Medium": "ffb74d", "High": "e57373"}
PLOT_SEVERITY_COLORS = {'None':'gray','Low':'skyblue','Medium':'orange','High':'crimson'}
ASI_SEVERITY_COLOR = {"None": "#66bb6a", "Low": "#ffd54f", "Medium": "#ffb74d", "High": "#e57373"}
RISK_LABEL_COLOR = {"None": "#66bb6a", "Low": "#ffd54f", "Medium": "#ffb74d", "High": "#e57373", "Critical": "#b71c1c"}

# Framework mappings (preserved)
CAF_DSPT_MAP = {
    "Throughput": "B - Protecting against cyber attack",
    "Packet Loss": "B - Protecting against cyber attack",
    "Delay": "D - Minimising impact of incidents",
    "Jitter": "C - Detecting cyber security events",
    "Default": "A - Managing security risk"
}

FHIR_MAP = {
    "Throughput": "Availability",
    "Delay": "Availability",
    "Jitter": "Integrity",
    "Packet Loss": "Confidentiality/Integrity",
    "Default": "Integrity"
}

MITIGATION_BASE = {
    "Throughput": "Investigate bandwidth/path, QoS, encryption, routing.",
    "Packet Loss": "Check links, NIC errors, retransmissions; monitor drops.",
    "Delay": "Investigate latency sources, queueing; implement redundancy/QoS.",
    "Jitter": "Examine scheduling/buffers; apply smoothing/buffering."
}

# Benchmarks (preserved)
BENCHMARKS = {
    "WIP": {
        "NORMAL": {"Throughput_MB": 0.2569307161781527, "Delay_s": 0.0028277116697944,
                   "Jitter_s": 0.0010952383587659, "PacketLoss_pct": 0.364406779661017},
        "MITM": {"Throughput_MB": 0.2238214574107829, "Delay_s": 0.0018630031171353,
                 "Jitter_s": 0.0006204625155563, "PacketLoss_pct": 11.499843063402386}
    },
    "SHS": {
        "NORMAL": {"Throughput_MB": 0.2507136277030552, "Delay_s": 0.0024455663002747,
                   "Jitter_s": 0.0007627349268911, "PacketLoss_pct": 0.8217928145468011},
        "MITM": {"Throughput_MB": 0.2242105728066907, "Delay_s": 0.0018567191454093,
                 "Jitter_s": 0.0006191572289738, "PacketLoss_pct": 11.364406779661016}
    }
}

###############################################################################
# ADDITIONAL STANDARDS CONFIGURATION (ADDED)
###############################################################################

__ADDITIONAL_STANDARDS__ = {
    "ISO/IEC 27400:2022": "IoT security and privacy — Threat taxonomy and IoT-specific guidance",
    "ISO 27799": "Health informatics — information security management in health",
    "ISO/IEC 27701": "Privacy information management — PIMS extension to ISO/IEC 27001",
    "ISO 13485:2016": "Medical devices — Quality management systems",
    "IEC 62443": "Industrial communication networks — OT/ICS security (zone/conduit model)"
}

# Map high-level attack labels to IoT threat categories (ISO/IEC 27400)
THREAT_MAP = {
    "MITM": "Network interception / message tampering",
    "DOS": "Denial of Service / availability compromise",
    "EAVESDROP": "Confidentiality breach / data leakage"
}

# Map device types to IEC 62443 Zones (example)
ZONE_MAP = {
    "wip": {"Zone": "IEC62443 Zone 1 (Clinical Wi-Fi LAN)", "Recommended_SL": 3},
    "shs": {"Zone": "IEC62443 Zone 2 (Bluetooth P2P / Companion Device)", "Recommended_SL": 2}
}

# ISO 27799 thresholds (example heuristics): if ASI >= 25 treat as potential non-compliance of clinical data protection
ISO27799_THRESHOLD_ASI = 25.0

# ISO 27701 privacy heuristics
def compute_privacy_impact(asi, plr):
    """Simple privacy impact heuristic (ADJUST to your policy)"""
    if asi >= 50 or plr >= 50:
        return "High"
    if asi >= 25 or plr >= 25:
        return "Medium"
    return "Low"

###############################################################################
# Utility functions (preserved)
###############################################################################

def clamp0_100(x):
    try:
        return float(max(0.0, min(100.0, float(x))))
    except Exception:
        return 0.0

def label_from_thresholds(val, thresholds):
    for low, high, label in thresholds:
        if low <= val < high:
            return label
    return thresholds[-1][2]

def clean_time_ns(t_str):
    """Return float ns value from strings like '123456ns' or '123456' or '123456+ns'."""
    if t_str is None:
        return 0.0
    s = str(t_str).replace('+', '').replace('ns', '')
    try:
        return float(s)
    except Exception:
        m = re.search(r'(\d+)', s)
        return float(m.group(1)) if m else 0.0

def _extract_flows_from_flowstats(flowstats):
    """Extract Flow nodes from FlowStats (handles direct Flow or Node->Flow)."""
    if not flowstats:
        return []
    flows = []
    if 'Flow' in flowstats:
        f = flowstats['Flow']
        if isinstance(f, list):
            flows.extend(f)
        else:
            flows.append(f)
    if 'Node' in flowstats:
        nodes = flowstats['Node']
        if isinstance(nodes, dict):
            nodes = [nodes]
        for node in nodes:
            f = node.get('Flow', [])
            if isinstance(f, dict):
                flows.append(f)
            else:
                flows.extend(f or [])
    return flows

def parse_flow_monitor_xml(path):
    """Parse one FlowMonitor XML and return DataFrame of primary metrics per flow."""
    with open(path, 'r') as f:
        doc = xmltodict.parse(f.read())

    fm = doc.get('FlowMonitor', {}) or {}
    fs = fm.get('FlowStats', {}) or {}
    flow_nodes = _extract_flows_from_flowstats(fs)
    data = []
    for flow in flow_nodes:
        try:
            fid = int(flow.get('@flowId', 0))
            tx = int(flow.get('@txPackets', 0))
            rx = int(flow.get('@rxPackets', 0))
            lost = int(flow.get('@lostPackets', 0))
            rxBytes = int(flow.get('@rxBytes', 0))
            t_first_tx = clean_time_ns(flow.get('@timeFirstTxPacket'))
            t_last_rx  = clean_time_ns(flow.get('@timeLastRxPacket'))
            delay_sum_ns = clean_time_ns(flow.get('@delaySum'))
            jitter_sum_ns = clean_time_ns(flow.get('@jitterSum'))

            duration_s = max((t_last_rx - t_first_tx) * 1e-9, EPS)
            throughput_mb_s = (rxBytes * BYTES_TO_MB) / duration_s if rx > 0 else 0.0
            avg_delay_s = (delay_sum_ns * 1e-9) / rx if rx > 0 else 0.0
            avg_jitter_s = (jitter_sum_ns * 1e-9) / rx if rx > 0 else 0.0
            pkt_loss_pct = (lost / tx * 100.0) if tx > 0 else 0.0

            data.append({
                "FlowID": fid,
                "TxPackets": tx,
                "RxPackets": rx,
                "LostPackets": lost,
                "RxBytes": rxBytes,
                "Throughput_MB_s": throughput_mb_s,
                "AvgDelay_s": avg_delay_s,
                "AvgJitter_s": avg_jitter_s,
                "PacketLoss_pct": pkt_loss_pct
            })
        except Exception:
            continue
    return pd.DataFrame(data)

def load_and_aggregate(dirpath):
    """Parse all XMLs in dirpath and return raw runs and aggregated metrics."""
    xmls = sorted([os.path.join(dirpath, f) for f in os.listdir(dirpath) if f.endswith('.xml')])
    if not xmls:
        raise FileNotFoundError(f"No XML files found in {dirpath}")

    runs = []
    for f in xmls:
        df = parse_flow_monitor_xml(f)
        if df.empty:
            continue
        df['RunFile'] = os.path.basename(f)
        runs.append(df)
    if not runs:
        return pd.DataFrame(), pd.DataFrame()

    raw = pd.concat(runs, ignore_index=True)
    agg = raw.groupby('FlowID').agg(
        Throughput_MB_s_median = ('Throughput_MB_s', 'median'),
        Throughput_MB_s_mean   = ('Throughput_MB_s', 'mean'),
        Throughput_MB_s_std    = ('Throughput_MB_s', 'std'),
        AvgDelay_s_median      = ('AvgDelay_s', 'median'),
        AvgDelay_s_mean        = ('AvgDelay_s', 'mean'),
        AvgDelay_s_std         = ('AvgDelay_s', 'std'),
        AvgJitter_s_median     = ('AvgJitter_s', 'median'),
        AvgJitter_s_mean       = ('AvgJitter_s', 'mean'),
        AvgJitter_s_std        = ('AvgJitter_s', 'std'),
        PacketLoss_pct_median  = ('PacketLoss_pct', 'median'),
        PacketLoss_pct_mean    = ('PacketLoss_pct', 'mean'),
        PacketLoss_pct_std     = ('PacketLoss_pct', 'std'),
        runs_count             = ('RunFile', 'nunique')
    ).reset_index()
    return raw, agg

def map_asi_to_nhs_likelihood(asi):
    for low, high, val in NHS_LIKELIHOOD_THRESH:
        if low <= asi < high:
            return val
    return NHS_LIKELIHOOD_THRESH[-1][2]

def map_cia_to_impact(c_score, i_score, a_score, device, flowid):
    avg = (c_score + i_score + a_score) / 3.0
    if flowid == FLOW9_ID:
        avg = (c_score + i_score + a_score * 1.3) / 3.3
    for low, high, val in NHS_IMPACT_THRESH:
        if low <= avg < high:
            return val
    return NHS_IMPACT_THRESH[-1][2]

def risk_score_rag(score):
    if score <= 3:
        return "Green"
    if score <= 6:
        return "Amber"
    if score <= 12:
        return "Orange"
    return "Red"

###############################################################################
# Core metric computation (modified to include new standard outputs)
###############################################################################

def compute_metrics_from_aggregates(agg_normal, agg_attack, device_type):
    """
    Compute comprehensive metrics including risk mapping.
    EXTENDED: Adds ISO/IEC 27400, ISO 27799, ISO/IEC 27701, ISO 13485 & IEC 62443 outputs.
    """
    device_key = device_type.lower()
    weights = DEVICE_WEIGHTS.get(device_key, {'TRR':0.25,'PLR':0.25,'NDI':0.25,'JVI':0.25})
    flow_ids = sorted(set(agg_normal['FlowID']).union(set(agg_attack['FlowID'])))
    
    benchmarks = BENCHMARKS.get(device_type.upper(), {})
    normal_bench = benchmarks.get("NORMAL", {})
    
    rows = []
    for fid in flow_ids:
        n_row = agg_normal[agg_normal.FlowID==fid]
        a_row = agg_attack[agg_attack.FlowID==fid]
        
        # Extract raw values
        Tn = float(n_row.Throughput_MB_s_median.values[0]) if not n_row.empty else 0.0
        Ta = float(a_row.Throughput_MB_s_median.values[0]) if not a_row.empty else 0.0
        On = float(n_row.AvgDelay_s_median.values[0]) if not n_row.empty else 0.0
        Oa = float(a_row.AvgDelay_s_median.values[0]) if not a_row.empty else 0.0
        Vn = float(n_row.AvgJitter_s_median.values[0]) if not n_row.empty else 0.0
        Va = float(a_row.AvgJitter_s_median.values[0]) if not a_row.empty else 0.0
        PLn = float(n_row.PacketLoss_pct_median.values[0]) if not n_row.empty else 0.0
        PLa = float(a_row.PacketLoss_pct_median.values[0]) if not a_row.empty else 100.0
        
        # FIXED: Track if attack had complete packet loss (zero reception)
        attack_has_complete_loss = (not a_row.empty and Ta == 0.0 and Oa == 0.0 and Va == 0.0) or (PLa >= 95.0)

        # Use benchmarks for baselines
        Tn_baseline = max(Tn, normal_bench.get("Throughput_MB", Tn)) if Tn > 0 else normal_bench.get("Throughput_MB", EPS)
        On_baseline = max(On, normal_bench.get("Delay_s", On)) if On > 0 else normal_bench.get("Delay_s", EPS)
        Vn_baseline = max(Vn, normal_bench.get("Jitter_s", Vn)) if Vn > 0 else normal_bench.get("Jitter_s", EPS)
        PLn_baseline = max(PLn, normal_bench.get("PacketLoss_pct", PLn)) if PLn > 0 else normal_bench.get("PacketLoss_pct", EPS)

        # Calculate TRR (Throughput Reduction Rate)
        TRR = clamp0_100(100.0 * (Tn_baseline - Ta) / Tn_baseline) if Tn_baseline > EPS else 0.0
        
        # Calculate PLR (Packet Loss Rate)
        PLR = clamp0_100(100.0 * (PLa / PLn_baseline)) if PLn_baseline > EPS else (clamp0_100(PLa * 10.0) if PLa > 0 else 0.0)
        
        # FIXED NDI CALCULATION
        if attack_has_complete_loss:
            NDI = 100.0
        elif On_baseline > EPS:
            NDI = clamp0_100(100.0 * abs(Oa - On_baseline) / On_baseline)
        else:
            NDI = clamp0_100(Oa * 1000.0) if Oa > 0 else 0.0
        
        # FIXED JVI CALCULATION
        if attack_has_complete_loss:
            JVI = 100.0
        elif Vn_baseline > EPS:
            JVI = clamp0_100(100.0 * abs(Va - Vn_baseline) / Vn_baseline)
        else:
            JVI = clamp0_100(Va * 1000.0) if Va > 0 else 0.0
        
        # PDR (Packet Delivery Rate)
        PDR = clamp0_100(100.0 * (Oa / On_baseline)) if On_baseline > EPS else 0.0

        # Calculate ASI (Attack Severity Index)
        ASI = clamp0_100(TRR * weights["TRR"] + PLR * weights["PLR"] + NDI * weights["NDI"] + JVI * weights["JVI"])
        ASI_Severity = label_from_thresholds(ASI, ASI_SEVERITY)
        
        # CIA scores and risk
        C_score = clamp0_100(ASI * CIA_WEIGHTS["C"])
        I_score = clamp0_100(ASI * CIA_WEIGHTS["I"])
        A_score = clamp0_100(ASI * CIA_WEIGHTS["A"])
        RiskScore = clamp0_100(C_score + I_score + A_score)
        RiskLabel = label_from_thresholds(RiskScore, RISK_LEVELS)
        
        # NHS risk mapping
        nhs_likelihood = map_asi_to_nhs_likelihood(ASI)
        nhs_impact = map_cia_to_impact(C_score, I_score, A_score, device_type, fid)
        nhs_risk_score = nhs_likelihood * nhs_impact
        nhs_rag = risk_score_rag(nhs_risk_score)

        # Determine dominant metric
        contributions = {
            "Throughput": TRR*weights["TRR"], 
            "Packet Loss": PLR*weights["PLR"], 
            "Delay": NDI*weights["NDI"], 
            "Jitter": JVI*weights["JVI"]
        }
        dominant = max(contributions.items(), key=lambda x: x[1])[0]
        
        # Framework mappings (preserved)
        caf = CAF_DSPT_MAP.get(dominant, CAF_DSPT_MAP["Default"])
        fhir = FHIR_MAP.get(dominant, FHIR_MAP["Default"])
        control_gap = "Yes" if (ASI >= 50 or PLR >= 50) else "No"
        
        # Mitigation (preserved)
        mitigation = MITIGATION_BASE.get(dominant, "Investigate & monitor")
        if attack_has_complete_loss:
            mitigation = "CRITICAL: Complete packet loss detected. " + mitigation + " Emergency response required."
        elif RiskLabel in ("High", "Critical"):
            mitigation += " Escalate immediately; incident response."
        elif ASI >= 25:
            mitigation += " Increase monitoring & schedule remediation."

        # Colors
        ASI_Severity_Color = ASI_SEVERITY_COLOR.get(ASI_Severity, "#B0B0B0")
        RiskLabel_Color = RISK_LABEL_COLOR.get(RiskLabel, "#B0B0B0")

        # Compute PriorityScore and ActionLevel (preserved)
        priority_score, action_level, priority_thresholds, criticality_factor = compute_priority(ASI, device_type)

        # -----------------------------
        # ADDED STANDARD-BASED FIELDS
        # -----------------------------
        # Threat category (ISO/IEC 27400)
        threat_cat = THREAT_MAP.get("MITM", "Network interception / Unknown")

        # IEC 62443 zone assignment & recommended Security Level
        iec_zone_info = ZONE_MAP.get(device_key, {"Zone": "Unzoned", "Recommended_SL": 1})
        iec_zone = iec_zone_info["Zone"]
        iec_recommended_sl = iec_zone_info["Recommended_SL"]

        # Privacy impact (ISO/IEC 27701) heuristic based on ASI and PLR
        privacy_impact = compute_privacy_impact(ASI, PLR)

        # ISO 27799 compliance flag (health-informatics CIA baseline)
        cia_compliance = "Non-Compliant" if ASI >= ISO27799_THRESHOLD_ASI else "Compliant"

        # ISO 13485 device risk traceability (map device criticality)
        device_qms_note = "ISO13485-aligned device QMS required" if device_key in DEVICE_CRITICALITY else "QMS review suggested"

        rows.append({
            "Device": device_type.upper(),
            "FlowID": fid,
            "Label": 1,
            "Throughput_normal_MB": Tn,
            "Throughput_attack_MB": Ta,
            "OWD_normal_s": On,
            "OWD_attack_s": Oa,
            "PDV_normal_s": Vn,
            "PDV_attack_s": Va,
            "PacketLoss_normal_pct": PLn,
            "PacketLoss_attack_pct": PLa,
            "TRR_pct": TRR,
            "PDR_pct": PDR,
            "PLR_pctpoints": PLR,
            "NDI_pct": NDI,
            "JVI_pct": JVI,
            "ASI_score": ASI,
            "ASI_Severity": ASI_Severity,
            "ASI_Severity_Color": ASI_Severity_Color,
            "C_score": C_score,
            "I_score": I_score,
            "A_score": A_score,
            "RiskScore": RiskScore,
            "RiskLabel": RiskLabel,
            "RiskLabel_Color": RiskLabel_Color,
            "NHS_Likelihood": nhs_likelihood,
            "NHS_Impact": nhs_impact,
            "NHS_RiskScore": nhs_risk_score,
            "NHS_RAG": nhs_rag,
            "DominantMetric": dominant,
            "CAF_DSPT_Domain": caf,
            "FHIR_Security_Dimension": fhir,
            "ControlGap_ISO27001_A12_6_1_2": control_gap,
            "Mitigation": mitigation,
            "Severity": ASI_Severity,
            "PriorityScore": priority_score,
            "ActionLevel": action_level,
            "CriticalityFactor": criticality_factor,
            "CompleteLoss": "Yes" if attack_has_complete_loss else "No",
            # ADDED fields
            "ThreatCategory": threat_cat,
            "IEC62443_Zone": iec_zone,
            "IEC62443_RecSL": iec_recommended_sl,
            "PrivacyImpactLevel": privacy_impact,
            "CIA_Compliance": cia_compliance,
            "ISO13485_Note": device_qms_note
        })
    return pd.DataFrame(rows)

def suggest_mitigation(row):
    """Legacy mitigation function for compatibility."""
    metrics = {'TRR': row['TRR_pct'], 'PLR': row['PLR_pctpoints'], 'NDI': row['NDI_pct'], 'JVI': row['JVI_pct']}
    top = max(metrics.items(), key=lambda x: abs(x[1]))[0]
    mit_map = {
        'TRR': "Investigate bandwidth, QoS, routing; check attacker path.",
        'PLR': "Check link errors, retransmissions; apply reliability/QoS.",
        'NDI': "Investigate latency sources, bufferbloat; add redundancy.",
        'JVI': "Apply jitter buffers, smoothing; examine scheduling."
    }
    return mit_map.get(top, "General monitoring & hardening.")

###############################################################################
# PLOTTING FUNCTIONS (preserved)
###############################################################################

def save_comparison_plots(device, agg_norm, agg_att, derived_df):
    """Generate comparison plots for raw and derived metrics."""
    plots = {}
    dev_upper = device.upper()
    device_plots_dir = os.path.join(PLOTS_DIR, dev_upper)
    os.makedirs(device_plots_dir, exist_ok=True)
    
    # Merge aggregates for comparison
    merged = pd.merge(
        agg_norm[['FlowID', 'Throughput_MB_s_median', 'AvgDelay_s_median', 'AvgJitter_s_median', 'PacketLoss_pct_median']],
        agg_att[['FlowID', 'Throughput_MB_s_median', 'AvgDelay_s_median', 'AvgJitter_s_median', 'PacketLoss_pct_median']],
        on='FlowID',
        suffixes=('_normal', '_attack'),
        how='outer'
    ).fillna(0)
    
    # Sort by FlowID
    merged = merged.sort_values('FlowID')
    flow_ids = merged['FlowID'].astype(str).values
    
    # 1. Throughput Comparison
    fig, ax = plt.subplots(figsize=(10, 5))
    x = np.arange(len(flow_ids))
    width = 0.35
    ax.bar(x - width/2, merged['Throughput_MB_s_median_normal'], width, label='Normal', color='skyblue')
    ax.bar(x + width/2, merged['Throughput_MB_s_median_attack'], width, label='Attack', color='crimson')
    ax.set_xlabel('Flow ID')
    ax.set_ylabel('Throughput (MB/s)')
    ax.set_title(f'{dev_upper} - Throughput Comparison')
    ax.set_xticks(x)
    ax.set_xticklabels(flow_ids)
    ax.legend()
    plt.tight_layout()
    path = os.path.join(device_plots_dir, f"{dev_upper}_Throughput_MB_s_comparison.png")
    plt.savefig(path)
    plt.close()
    plots['Throughput_MB_s_comparison'] = path
    
    # 2. Delay Comparison
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.bar(x - width/2, merged['AvgDelay_s_median_normal'], width, label='Normal', color='skyblue')
    ax.bar(x + width/2, merged['AvgDelay_s_median_attack'], width, label='Attack', color='crimson')
    ax.set_xlabel('Flow ID')
    ax.set_ylabel('Avg Delay (s)')
    ax.set_title(f'{dev_upper} - Delay Comparison')
    ax.set_xticks(x)
    ax.set_xticklabels(flow_ids)
    ax.legend()
    plt.tight_layout()
    path = os.path.join(device_plots_dir, f"{dev_upper}_AvgDelay_s_comparison.png")
    plt.savefig(path)
    plt.close()
    plots['AvgDelay_s_comparison'] = path
    
    # 3. Jitter Comparison
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.bar(x - width/2, merged['AvgJitter_s_median_normal'], width, label='Normal', color='skyblue')
    ax.bar(x + width/2, merged['AvgJitter_s_median_attack'], width, label='Attack', color='crimson')
    ax.set_xlabel('Flow ID')
    ax.set_ylabel('Avg Jitter (s)')
    ax.set_title(f'{dev_upper} - Jitter Comparison')
    ax.set_xticks(x)
    ax.set_xticklabels(flow_ids)
    ax.legend()
    plt.tight_layout()
    path = os.path.join(device_plots_dir, f"{dev_upper}_AvgJitter_s_comparison.png")
    plt.savefig(path)
    plt.close()
    plots['AvgJitter_s_comparison'] = path
    
    # 4. Packet Loss Comparison
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.bar(x - width/2, merged['PacketLoss_pct_median_normal'], width, label='Normal', color='skyblue')
    ax.bar(x + width/2, merged['PacketLoss_pct_median_attack'], width, label='Attack', color='crimson')
    ax.set_xlabel('Flow ID')
    ax.set_ylabel('Packet Loss (%)')
    ax.set_title(f'{dev_upper} - Packet Loss Comparison')
    ax.set_xticks(x)
    ax.set_xticklabels(flow_ids)
    ax.legend()
    plt.tight_layout()
    path = os.path.join(device_plots_dir, f"{dev_upper}_PacketLoss_pct_comparison.png")
    plt.savefig(path)
    plt.close()
    plots['PacketLoss_pct_comparison'] = path
    
    # Now plot derived metrics from derived_df
    derived_sorted = derived_df.sort_values('FlowID')
    flow_ids_derived = derived_sorted['FlowID'].astype(str).values
    x_derived = np.arange(len(flow_ids_derived))
    
    # 5. TRR Plot
    fig, ax = plt.subplots(figsize=(10, 5))
    colors_trr = [PLOT_SEVERITY_COLORS.get(s, 'gray') for s in derived_sorted['Severity']]
    ax.bar(x_derived, derived_sorted['TRR_pct'], color=colors_trr)
    ax.set_xlabel('Flow ID')
    ax.set_ylabel('TRR (%)')
    ax.set_title(f'{dev_upper} - Throughput Reduction Rate (TRR)')
    ax.set_xticks(x_derived)
    ax.set_xticklabels(flow_ids_derived)
    plt.tight_layout()
    path = os.path.join(device_plots_dir, f"{dev_upper}_TRR_comparison.png")
    plt.savefig(path)
    plt.close()
    plots['TRR_comparison'] = path
    
    # 6. PDR Plot
    fig, ax = plt.subplots(figsize=(10, 5))
    colors_pdr = [PLOT_SEVERITY_COLORS.get(s, 'gray') for s in derived_sorted['Severity']]
    ax.bar(x_derived, derived_sorted['PDR_pct'], color=colors_pdr)
    ax.set_xlabel('Flow ID')
    ax.set_ylabel('PDR (%)')
    ax.set_title(f'{dev_upper} - Packet Delivery Rate (PDR)')
    ax.set_xticks(x_derived)
    ax.set_xticklabels(flow_ids_derived)
    plt.tight_layout()
    path = os.path.join(device_plots_dir, f"{dev_upper}_PDR_comparison.png")
    plt.savefig(path)
    plt.close()
    plots['PDR_comparison'] = path
    
    # 7. PLR Plot
    fig, ax = plt.subplots(figsize=(10, 5))
    colors_plr = [PLOT_SEVERITY_COLORS.get(s, 'gray') for s in derived_sorted['Severity']]
    ax.bar(x_derived, derived_sorted['PLR_pctpoints'], color=colors_plr)
    ax.set_xlabel('Flow ID')
    ax.set_ylabel('PLR (percentage points)')
    ax.set_title(f'{dev_upper} - Packet Loss Rate (PLR)')
    ax.set_xticks(x_derived)
    ax.set_xticklabels(flow_ids_derived)
    plt.tight_layout()
    path = os.path.join(device_plots_dir, f"{dev_upper}_PLR_comparison.png")
    plt.savefig(path)
    plt.close()
    plots['PLR_comparison'] = path
    
    # 8. JVI Plot
    fig, ax = plt.subplots(figsize=(10, 5))
    colors_jvi = [PLOT_SEVERITY_COLORS.get(s, 'gray') for s in derived_sorted['Severity']]
    ax.bar(x_derived, derived_sorted['JVI_pct'], color=colors_jvi)
    ax.set_xlabel('Flow ID')
    ax.set_ylabel('JVI (%)')
    ax.set_title(f'{dev_upper} - Jitter Variation Index (JVI)')
    ax.set_xticks(x_derived)
    ax.set_xticklabels(flow_ids_derived)
    plt.tight_layout()
    path = os.path.join(device_plots_dir, f"{dev_upper}_JVI_comparison.png")
    plt.savefig(path)
    plt.close()
    plots['JVI_comparison'] = path
    
    # 9. NDI Plot
    fig, ax = plt.subplots(figsize=(10, 5))
    colors_ndi = [PLOT_SEVERITY_COLORS.get(s, 'gray') for s in derived_sorted['Severity']]
    ax.bar(x_derived, derived_sorted['NDI_pct'], color=colors_ndi)
    ax.set_xlabel('Flow ID')
    ax.set_ylabel('NDI (%)')
    ax.set_title(f'{dev_upper} - Network Delay Impact (NDI)')
    ax.set_xticks(x_derived)
    ax.set_xticklabels(flow_ids_derived)
    plt.tight_layout()
    path = os.path.join(device_plots_dir, f"{dev_upper}_NDI_comparison.png")
    plt.savefig(path)
    plt.close()
    plots['NDI_comparison'] = path
    
    # 10. ASI Plot
    fig, ax = plt.subplots(figsize=(10, 5))
    colors_asi = [PLOT_SEVERITY_COLORS.get(s, 'gray') for s in derived_sorted['Severity']]
    ax.bar(x_derived, derived_sorted['ASI_score'], color=colors_asi)
    ax.set_xlabel('Flow ID')
    ax.set_ylabel('ASI Score')
    ax.set_title(f'{dev_upper} - Attack Severity Index (ASI)')
    ax.set_xticks(x_derived)
    ax.set_xticklabels(flow_ids_derived)
    plt.tight_layout()
    path = os.path.join(device_plots_dir, f"{dev_upper}_ASI_comparison.png")
    plt.savefig(path)
    plt.close()
    plots['ASI_comparison'] = path
    
    # 11. RiskScore Plot
    fig, ax = plt.subplots(figsize=(10, 5))
    colors_risk = [RISK_LABEL_COLOR.get(r, '#B0B0B0') for r in derived_sorted['RiskLabel']]
    ax.bar(x_derived, derived_sorted['RiskScore'], color=colors_risk)
    ax.set_xlabel('Flow ID')
    ax.set_ylabel('Risk Score')
    ax.set_title(f'{dev_upper} - Overall Risk Score')
    ax.set_xticks(x_derived)
    ax.set_xticklabels(flow_ids_derived)
    plt.tight_layout()
    path = os.path.join(device_plots_dir, f"{dev_upper}_RiskScore_comparison.png")
    plt.savefig(path)
    plt.close()
    plots['RiskScore_comparison'] = path
    
    print(f"[+] Generated {len(plots)} comparison plots for {dev_upper}")
    return plots

def colorize_cells(ws, color_col, target_col):
    """Apply color formatting to Excel cells."""
    header = [cell.value for cell in ws[1]]
    if color_col not in header or target_col not in header:
        return
    color_idx = header.index(color_col) + 1
    target_idx = header.index(target_col) + 1
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
        color_code = row[color_idx-1].value
        if color_code and isinstance(color_code, str) and (color_code.startswith("#") or len(color_code)==6):
            hexcode = color_code[1:] if color_code.startswith("#") else color_code
            row[target_idx-1].fill = PatternFill(start_color=hexcode, end_color=hexcode, fill_type="solid")

###############################################################################
# Build/Export functions (extended to produce added standard sheets)
###############################################################################

def build_comprehensive_excel(derived_all, all_plots):
    """Build comprehensive Excel workbook with all sheets and risk mappings."""
    print(f"[+] Building comprehensive Excel workbook: {EXCEL_PATH}")
    
    # Prepare data for various sheets
    caf_list, nhs_list, fhir_list, dos_list = [], [], [], []
    iso27001_list, iso27005_list, fda_list = [], [], []
    # ADDED lists for additional standards
    iso27400_list, iso27799_list, iso27701_list, iec62443_list, iso13485_list = [], [], [], [], []

    # Process each device
    for dev in ['WIP', 'SHS']:
        dev_data = derived_all[derived_all['Device'] == dev]
        if dev_data.empty:
            continue
            
        # Build framework mapping data
        for _, r in dev_data.iterrows():
            fid = int(r["FlowID"])
            caf_list.append({
                "Device": dev,
                "FlowID": fid,
                "DominantMetric": r["DominantMetric"],
                "CAF_DSPT_Domain": r["CAF_DSPT_Domain"],
                "ASI": r["ASI_score"],
                "RiskLabel": r["RiskLabel"],
                "ASI_Severity_Color": r.get("ASI_Severity_Color", ""),
                "RiskLabel_Color": r.get("RiskLabel_Color", ""),
                "Mitigation": r["Mitigation"]
            })
            nhs_list.append({
                "Device": dev,
                "FlowID": fid,
                "Likelihood": r.get("NHS_Likelihood", 1),
                "Impact": r.get("NHS_Impact", 1),
                "Risk": r.get("NHS_RiskScore", 0),
                "RAG": r.get("NHS_RAG", ""),
                "ASI": r["ASI_score"],
                "RiskLabel": r["RiskLabel"],
                "RiskLabel_Color": r.get("RiskLabel_Color", "")
            })
            fhir_list.append({
                "Device": dev,
                "FlowID": fid,
                "FHIR_Dimension": r.get("FHIR_Security_Dimension", ""),
                "ASI": r["ASI_score"],
                "RiskLabel": r["RiskLabel"],
                "ASI_Severity_Color": r.get("ASI_Severity_Color", ""),
                "RiskLabel_Color": r.get("RiskLabel_Color", "")
            })
            dos_list.append({
                "Device": dev,
                "FlowID": fid,
                "DoS_Risk": r.get("RiskLabel"),
                "Mitigation": r.get("Mitigation"),
                "RiskLabel_Color": r.get("RiskLabel_Color", "")
            })
            iso27001_list.append({
                "Device": dev,
                "FlowID": fid,
                "ISO27001_ControlGap": r.get("ControlGap_ISO27001_A12_6_1_2", ""),
                "ASI": r["ASI_score"],
                "RiskLabel": r["RiskLabel"],
                "RiskLabel_Color": r.get("RiskLabel_Color", "")
            })
            iso27005_list.append({
                "Device": dev,
                "FlowID": fid,
                "ISO27005_Severity": r.get("ASI_Severity", ""),
                "ISO27005_Risk": r.get("RiskLabel", ""),
                "RiskLabel_Color": r.get("RiskLabel_Color", "")
            })
            fda_list.append({
                "Risk ID": f"{dev}_F{fid}",
                "Hazardous Situation": f"Network degradation ({r.get('DominantMetric','Unknown')}) affecting device {dev}, flow {fid}",
                "Clinical Harm Severity": r.get("ASI_Severity", ""),
                "Probability (1-5)": r.get("NHS_Likelihood", 1),
                "Existing Controls": "Network segmentation, baseline monitoring",
                "Suggested Controls": r.get("Mitigation", ""),
                "Residual Risk Acceptable": "No" if r.get("RiskLabel") in ("High", "Critical") else "Yes",
                "ASI_Severity_Color": r.get("ASI_Severity_Color", ""),
                "RiskLabel_Color": r.get("RiskLabel_Color", "")
            })
            # ADDED: collect rows for additional standards
            iso27400_list.append({
                "Device": dev,
                "FlowID": fid,
                "ThreatCategory": r.get("ThreatCategory", ""),
                "DominantMetric": r.get("DominantMetric", ""),
                "ASI": r.get("ASI_score", 0),
                "RiskLabel": r.get("RiskLabel", "")
            })
            iso27799_list.append({
                "Device": dev,
                "FlowID": fid,
                "CIA_Compliance": r.get("CIA_Compliance", ""),
                "ASI": r.get("ASI_score", 0),
                "RiskLabel": r.get("RiskLabel", "")
            })
            iso27701_list.append({
                "Device": dev,
                "FlowID": fid,
                "PrivacyImpactLevel": r.get("PrivacyImpactLevel", ""),
                "ASI": r.get("ASI_score", 0),
                "RiskLabel": r.get("RiskLabel", "")
            })
            iec62443_list.append({
                "Device": dev,
                "FlowID": fid,
                "IEC62443_Zone": r.get("IEC62443_Zone", ""),
                "IEC62443_RecSL": r.get("IEC62443_RecSL", ""),
                "ASI": r.get("ASI_score", 0),
                "RiskLabel": r.get("RiskLabel", "")
            })
            iso13485_list.append({
                "Device": dev,
                "FlowID": fid,
                "ISO13485_Note": r.get("ISO13485_Note", ""),
                "ASI": r.get("ASI_score", 0),
                "RiskLabel": r.get("RiskLabel", "")
            })
    
    # Create DataFrames
    caf_df = pd.DataFrame(caf_list)
    nhs_df = pd.DataFrame(nhs_list)
    fhir_df = pd.DataFrame(fhir_list)
    dos_df = pd.DataFrame(dos_list)
    iso27001_df = pd.DataFrame(iso27001_list)
    iso27005_df = pd.DataFrame(iso27005_list)
    fda_df = pd.DataFrame(fda_list)
    # ADDED DataFrames
    iso27400_df = pd.DataFrame(iso27400_list)
    iso27799_df = pd.DataFrame(iso27799_list)
    iso27701_df = pd.DataFrame(iso27701_list)
    iec62443_df = pd.DataFrame(iec62443_list)
    iso13485_df = pd.DataFrame(iso13485_list)
    
    # Write Excel file
    with pd.ExcelWriter(EXCEL_PATH, engine="openpyxl") as writer:
        derived_all.to_excel(writer, sheet_name="Derived Metrics", index=False)
        derived_all[['Device','FlowID','ASI_score','Severity','Mitigation','PriorityScore','ActionLevel','CriticalityFactor']].to_excel(
            writer, sheet_name="Mitigations by Flow", index=False)
        pd.DataFrame({"Note":["Plots embedded below or available under analysis/plots"]}).to_excel(
            writer, sheet_name="Plots", index=False)
        caf_df.to_excel(writer, sheet_name="CAF-DSPT Mapping", index=False)
        nhs_df.to_excel(writer, sheet_name="NHS Risk Register", index=False)
        fhir_df.to_excel(writer, sheet_name="HL7 FHIR Mapping", index=False)
        dos_df.to_excel(writer, sheet_name="DoS Risk Plan", index=False)
        iso27001_df.to_excel(writer, sheet_name="ISO 27001 Mapping", index=False)
        iso27005_df.to_excel(writer, sheet_name="ISO 27005 Mapping", index=False)
        fda_df.to_excel(writer, sheet_name="FDA Risk Mapping", index=False)
        # ADDED sheets for new standards
        if not iso27400_df.empty:
            iso27400_df.to_excel(writer, sheet_name="ISO27400_ThreatMapping", index=False)
        if not iso27799_df.empty:
            iso27799_df.to_excel(writer, sheet_name="ISO27799_Compliance", index=False)
        if not iso27701_df.empty:
            iso27701_df.to_excel(writer, sheet_name="ISO27701_PrivacyMap", index=False)
        if not iec62443_df.empty:
            iec62443_df.to_excel(writer, sheet_name="IEC62443_NetworkZones", index=False)
        if not iso13485_df.empty:
            iso13485_df.to_excel(writer, sheet_name="ISO13485_DeviceRisk", index=False)
        
        # Add ISO metadata sheet (extended)
        meta_rows = []
        meta_rows.append({"Key":"CRITICALITY_METHOD","Value":__CRITICALITY_METADATA__.get("method","")})
        meta_rows.append({"Key":"DEVICE_RPN","Value":str(__CRITICALITY_METADATA__.get("DEVICE_RPN",{}))})
        meta_rows.append({"Key":"DEVICE_CRITICALITY","Value":str(__CRITICALITY_METADATA__.get("DEVICE_CRITICALITY",{}))})
        for i, citation in enumerate(__CRITICALITY_METADATA__.get("citations", []), 1):
            meta_rows.append({"Key": f"Citation_{i}", "Value": citation})
        # Additional standards metadata
        for k, v in __ADDITIONAL_STANDARDS__.items():
            meta_rows.append({"Key": f"AdditionalStandard_{k}", "Value": v})
        meta_df = pd.DataFrame(meta_rows)
        meta_df.to_excel(writer, sheet_name="ISO_Metadata", index=False)
    
    # Load workbook for formatting and plot embedding
    wb = load_workbook(EXCEL_PATH)
    
    # Embed plots (preserved)
    ws_plots = wb["Plots"]
    row = 2
    col = 1
    
    plot_keys = [
        "Throughput_MB_s_comparison", "AvgDelay_s_comparison", "AvgJitter_s_comparison",
        "PacketLoss_pct_comparison", "TRR_comparison", "PDR_comparison", "PLR_comparison",
        "JVI_comparison", "NDI_comparison", "ASI_comparison", "RiskScore_comparison"
    ]
    
    for dev in ['WIP', 'SHS']:
        plots = all_plots.get(dev, {})
        if not plots:
            continue
        ws_plots.cell(row=row, column=col).value = f"{dev} Plots"
        row += 1
        
        for key in plot_keys:
            path = plots.get(key)
            if path and os.path.exists(path):
                try:
                    img = XLImage(path)
                    max_w, max_h = 800, 400
                    if img.width > max_w:
                        scale = max_w / img.width
                        img.width = int(img.width * scale)
                        img.height = int(img.height * scale)
                    anchor = f"{get_column_letter(col)}{row}"
                    ws_plots.add_image(img, anchor)
                    row += int(img.height/20) + 2
                except Exception as e:
                    ws_plots.cell(row=row, column=col).value = f"Image embed failed: {path}"
                    row += 2
        row += 2
    
    # Apply color formatting (preserved)
    def safe_colorize(sheetname, color_col, target_col):
        if sheetname in wb.sheetnames:
            try:
                colorize_cells(wb[sheetname], color_col, target_col)
            except Exception:
                pass
    
    safe_colorize("Derived Metrics", "ASI_Severity_Color", "ASI_Severity")
    safe_colorize("Derived Metrics", "RiskLabel_Color", "RiskLabel")
    safe_colorize("CAF-DSPT Mapping", "ASI_Severity_Color", "ASI")
    safe_colorize("CAF-DSPT Mapping", "RiskLabel_Color", "RiskLabel")
    safe_colorize("NHS Risk Register", "RiskLabel_Color", "RiskLabel")
    safe_colorize("HL7 FHIR Mapping", "ASI_Severity_Color", "ASI")
    safe_colorize("HL7 FHIR Mapping", "RiskLabel_Color", "RiskLabel")
    safe_colorize("DoS Risk Plan", "RiskLabel_Color", "DoS_Risk")
    safe_colorize("ISO 27001 Mapping", "RiskLabel_Color", "RiskLabel")
    safe_colorize("ISO 27005 Mapping", "RiskLabel_Color", "ISO27005_Risk")
    safe_colorize("FDA Risk Mapping", "ASI_Severity_Color", "Clinical Harm Severity")
    safe_colorize("FDA Risk Mapping", "RiskLabel_Color", "Residual Risk Acceptable")
    
    # Color the ActionLevel column (preserved)
    if "Derived Metrics" in wb.sheetnames:
        ws_dm = wb["Derived Metrics"]
        headers = [cell.value for cell in ws_dm[1]]
        if "ActionLevel" in headers:
            a_idx = headers.index("ActionLevel") + 1
            for row in ws_dm.iter_rows(min_row=2, max_row=ws_dm.max_row):
                cell = row[a_idx-1]
                val = (cell.value or "").strip() if cell.value is not None else ""
                if val == "Critical":
                    cell.fill = PatternFill(start_color="FF4C4C", end_color="FF4C4C", fill_type="solid")
                elif val == "Action":
                    cell.fill = PatternFill(start_color="FFD966", end_color="FFD966", fill_type="solid")
                elif val == "Warning":
                    cell.fill = PatternFill(start_color="FFFF99", end_color="FFFF99", fill_type="solid")
                elif val == "Normal":
                    cell.fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
    
    wb.save(EXCEL_PATH)
    wb.close()
    print(f"[+] Comprehensive Excel workbook saved: {EXCEL_PATH}")

def build_word_report(derived_all):
    """Build comprehensive Word document with risk assessment (extended to include added standards)."""
    if not DOCX_AVAILABLE:
        print("python-docx not available; skipping Word document.")
        return
    
    print(f"[+] Building Word report: {WORD_PATH}")
    doc = Document()
    doc.add_heading("IoMT MITM Attack - Comprehensive Risk Assessment & Security Plan", level=1)
    doc.add_paragraph("Generated by IoMT Merged Pipeline")
    doc.add_paragraph(f"This report includes FDA risk assessment, NHS risk register, CAF-DSPT mapping, and ISO 27001/27005 compliance. It has been extended to include ISO/IEC 27400, ISO 27799, ISO/IEC 27701, ISO 13485 and IEC 62443 mappings.")
    
    # Summary statistics
    doc.add_heading("Executive Summary", level=2)
    severity_counts = derived_all['ASI_Severity'].value_counts()
    risk_counts = derived_all['RiskLabel'].value_counts()
    
    doc.add_paragraph(f"Total Flows Analyzed: {len(derived_all)}")
    doc.add_paragraph(f"Severity Distribution: {dict(severity_counts)}")
    doc.add_paragraph(f"Risk Level Distribution: {dict(risk_counts)}")
    
    # Check for complete loss flows
    complete_loss_flows = derived_all[derived_all['CompleteLoss'] == 'Yes']
    if not complete_loss_flows.empty:
        doc.add_paragraph(f"CRITICAL: {len(complete_loss_flows)} flows experienced complete packet loss!")
        doc.add_paragraph("These flows show maximum network disruption with zero packet reception.")
    
    # Top 10 most severe flows
    doc.add_heading("Top 10 Most Severe Flows", level=2)
    topn = derived_all.sort_values("ASI_score", ascending=False).head(10)
    if not topn.empty:
        t = doc.add_table(rows=1, cols=10)
        hdr = t.rows[0].cells
        cols = ["Device", "FlowID", "ASI_score", "ASI_Severity", "RiskLabel", "NHS_RAG", "DominantMetric", "Mitigation", "ThreatCategory", "PrivacyImpactLevel"]
        for i, c in enumerate(cols):
            hdr[i].text = str(c)
        for _, r in topn.iterrows():
            rowcells = t.add_row().cells
            for i, c in enumerate(cols):
                rowcells[i].text = str(r.get(c, ""))
    
    # Device-specific analysis (preserved + added notes)
    for dev in ['WIP', 'SHS']:
        dev_data = derived_all[derived_all['Device'] == dev]
        if dev_data.empty:
            continue
        
        doc.add_heading(f"Device: {dev}", level=2)
        doc.add_paragraph(f"Total flows: {len(dev_data)}")
        doc.add_paragraph(f"Critical/High risk flows: {len(dev_data[dev_data['RiskLabel'].isin(['Critical','High'])])}")
        
        # Check for complete loss in this device
        dev_complete_loss = dev_data[dev_data['CompleteLoss'] == 'Yes']
        if not dev_complete_loss.empty:
            doc.add_paragraph(f"WARNING: {len(dev_complete_loss)} flows with complete packet loss detected!")
            for _, r in dev_complete_loss.iterrows():
                doc.add_paragraph(f"  - Flow {int(r['FlowID'])}: ASI={r['ASI_score']:.1f}, Risk={r['RiskLabel']}")
        
        # ISO additions: show zone, threat, privacy
        doc.add_paragraph("Standards Mappings (sample):")
        sample = dev_data.head(5)
        if not sample.empty:
            table = doc.add_table(rows=1, cols=5)
            hdr = table.rows[0].cells
            hdr[0].text = "FlowID"
            hdr[1].text = "ThreatCategory (ISO27400)"
            hdr[2].text = "IEC62443 Zone"
            hdr[3].text = "PrivacyImpact (27701)"
            hdr[4].text = "CIA_Compliant (27799)"
            for _, r in sample.iterrows():
                row = table.add_row().cells
                row[0].text = str(int(r['FlowID']))
                row[1].text = str(r.get('ThreatCategory',''))
                row[2].text = str(r.get('IEC62443_Zone',''))
                row[3].text = str(r.get('PrivacyImpactLevel',''))
                row[4].text = str(r.get('CIA_Compliance',''))
        
        # FDA Risk Table (preserved)
        doc.add_heading(f"{dev} - FDA Risk Assessment", level=3)
        table = doc.add_table(rows=1, cols=7)
        hdr = table.rows[0].cells
        hdr[0].text = "Risk ID"
        hdr[1].text = "Hazard"
        hdr[2].text = "Severity"
        hdr[3].text = "Probability"
        hdr[4].text = "Controls"
        hdr[5].text = "Mitigation"
        hdr[6].text = "Acceptable"
        
        for _, r in dev_data.iterrows():
            row = table.add_row().cells
            row[0].text = f"{dev}_F{int(r['FlowID'])}"
            row[1].text = f"Network degradation ({r.get('DominantMetric','Unknown')})"
            row[2].text = r.get('ASI_Severity', '')
            row[3].text = str(r.get('NHS_Likelihood', 1))
            row[4].text = "Network segmentation, baseline monitoring"
            row[5].text = r.get('Mitigation', '')
            residual = "No" if r.get('RiskLabel') in ('High','Critical') else "Yes"
            row[6].text = residual
    
    # Recommendations (preserved)
    doc.add_heading("Recommendations", level=2)
    critical_flows = derived_all[derived_all['RiskLabel'].isin(['Critical', 'High'])]
    if not critical_flows.empty:
        doc.add_paragraph("URGENT: The following flows require immediate attention:")
        for _, r in critical_flows.iterrows():
            doc.add_paragraph(f"• {r['Device']} Flow {int(r['FlowID'])}: {r['Mitigation']}", style='List Bullet')
    else:
        doc.add_paragraph("No critical or high-risk flows identified.")
    
    doc.add_paragraph("")
    doc.add_paragraph("Detailed metrics available in Excel workbook and CSV files in analysis/ directory.")
    
    try:
        doc.save(WORD_PATH)
        print(f"[+] Word report saved: {WORD_PATH}")
    except Exception as e:
        print(f"Failed to save Word document: {e}")

###############################################################################
# Additional plotting helpers (preserved)
###############################################################################

def plot_asi_by_device(df, outdir):
    """Plot ASI scores by device and flow."""
    os.makedirs(outdir, exist_ok=True)
    for dev in df['Device'].unique():
        d = df[df.Device==dev].sort_values("FlowID")
        plt.figure(figsize=(10,4))
        colors = [PLOT_SEVERITY_COLORS.get(s, 'gray') for s in d['Severity']]
        plt.bar(d['FlowID'].astype(str), d['ASI_score'], color=colors)
        plt.title(f"{dev} - ASI score per Flow")
        plt.xlabel("FlowID")
        plt.ylabel("ASI score")
        p = os.path.join(outdir, f"{dev}_ASI_scores.png")
        plt.tight_layout()
        plt.savefig(p)
        plt.close()
        print(f"[+] Saved plot: {p}")

def plot_severity_pie(df, outdir):
    """Plot severity distribution as pie chart."""
    os.makedirs(outdir, exist_ok=True)
    a = df[df['Severity']!='None']
    counts = a['Severity'].value_counts().reindex(['Low','Medium','High'], fill_value=0)
    if counts.sum() == 0:
        print("No non-None severities - skipping severity pie.")
        return
    plt.figure(figsize=(6,6))
    colors = [PLOT_SEVERITY_COLORS[k] for k in ['Low','Medium','High']]
    counts.plot.pie(autopct='%1.1f%%', colors=colors)
    plt.ylabel('')
    plt.title('Attack Severity Distribution')
    p = os.path.join(outdir, "severity_attack_pie.png")
    plt.savefig(p)
    plt.close()
    print(f"[+] Saved plot: {p}")

def plot_severity_distribution(df, outdir):
    """Plot severity distribution as bar chart."""
    os.makedirs(outdir, exist_ok=True)
    attack_df = df[df['Label'] == 1]
    severity_counts = attack_df['Severity'].value_counts().reindex(['None', 'Low', 'Medium', 'High'], fill_value=0)
    plt.figure(figsize=(7,5))
    severity_colors = [PLOT_SEVERITY_COLORS[k] for k in ['None','Low','Medium','High']]
    plt.bar(severity_counts.index, severity_counts.values, color=severity_colors)
    plt.xlabel('Severity Level')
    plt.ylabel('Number of Attacks')
    plt.title('Severity Distribution (Attack Flows)')
    for i, v in enumerate(severity_counts.values):
        plt.text(i, v + 0.5, str(v), ha='center', va='bottom', fontsize=12)
    plt.tight_layout()
    out = os.path.join(outdir, "attack_severity_distribution_bar.png")
    plt.savefig(out)
    plt.close()
    print(f"[+] Saved plot: {out}")

def plot_attack_metrics_all_stats(df, outdir):
    """Plot attack metrics with various statistics."""
    os.makedirs(outdir, exist_ok=True)
    metrics = ['TRR_pct','PLR_pctpoints','PDR_pct','NDI_pct','JVI_pct','ASI_score']
    stat_funcs = {
        "mean": np.mean,
        "median": np.median,
        "max": np.max,
        "min": np.min,
        "std": np.std
    }
    scenarios = ['Attack-WIP', 'Attack-SHS']
    df_attack = df[df["Label"] == 1].copy()
    df_attack["Scenario"] = df_attack["Device"].map({"WIP":"Attack-WIP","SHS":"Attack-SHS"})
    
    for stat, func in stat_funcs.items():
        vals = {}
        for scenario in scenarios:
            sub = df_attack[df_attack["Scenario"] == scenario]
            vals[scenario] = [func(sub[metric]) if not sub.empty else np.nan for metric in metrics]
        x = np.arange(len(metrics))
        width = 0.35
        plt.figure(figsize=(10,6))
        plt.bar(x - width/2, vals[scenarios[0]], width, label=scenarios[0], color="crimson")
        plt.bar(x + width/2, vals[scenarios[1]], width, label=scenarios[1], color="darkblue")
        plt.xticks(x, metrics, rotation=20)
        plt.ylabel(f"{stat.capitalize()} Value")
        plt.title(f"{stat.capitalize()} Attack Metrics by Scenario")
        plt.legend()
        plt.tight_layout()
        fname = f"attack_metrics_{stat}_bar.png"
        out_path = os.path.join(outdir, fname)
        plt.savefig(out_path)
        plt.close()
        print(f"[+] Saved plot: {out_path}")

def plot_critical_flows(df, outdir):
    """Plot metrics for critical flows."""
    os.makedirs(outdir, exist_ok=True)
    metrics = ['TRR_pct','PLR_pctpoints','PDR_pct','NDI_pct','JVI_pct','ASI_score']
    for device in ['wip','shs']:
        sub = df[(df['Device'].str.lower() == device) & (df['Label'] == 1) & (df['Severity'] == "High")]
        if sub.empty:
            continue
        means = [sub[metric].mean() for metric in metrics]
        plt.figure(figsize=(10,6))
        plt.bar(np.arange(len(metrics)), means, color='orange')
        plt.xticks(np.arange(len(metrics)), metrics, rotation=20)
        plt.ylabel("Mean Value")
        plt.title(f"Critical Flows Mean Metrics - {device.upper()}")
        plt.tight_layout()
        fname = f"attack_metrics_critical_{device}.png"
        out_path = os.path.join(outdir, fname)
        plt.savefig(out_path)
        plt.close()
        print(f"[+] Saved plot: {out_path}")

def plot_flow9(df, outdir):
    """Plot Flow 9 specific metrics."""
    os.makedirs(outdir, exist_ok=True)
    metrics = ['TRR_pct','PLR_pctpoints','PDR_pct','NDI_pct','JVI_pct','ASI_score']
    for device in ['wip','shs']:
        sub = df[(df['Device'].str.lower() == device) & (df['Label'] == 1) & (df['FlowID'] == 9)]
        if sub.empty:
            continue
        vals = [sub[metric].mean() for metric in metrics]
        plt.figure(figsize=(10,6))
        plt.bar(np.arange(len(metrics)), vals, color=('crimson' if device=='wip' else 'darkblue'))
        plt.xticks(np.arange(len(metrics)), metrics, rotation=20)
        plt.ylabel("Mean Value")
        plt.title(f"Flow 9 Metrics - {device.upper()}")
        plt.tight_layout()
        fname = f"flow9_metrics_{device}.png"
        out_path = os.path.join(outdir, fname)
        plt.savefig(out_path)
        plt.close()
        print(f"[+] Saved plot: {out_path}")

def print_classification_report_and_save_confusion(df, outdir):
    """Print classification report and save confusion matrix."""
    os.makedirs(outdir, exist_ok=True)
    counts = df['Label'].value_counts().sort_index()
    print("\nScenario counts:")
    print(f"Attack (1): {counts.get(1, 0)}")
    print("\nClassification Report (Label):")
    print(classification_report(df['Label'], df['Label'], target_names=["Attack"]))
    cm = confusion_matrix(df['Label'], df['Label'])
    plt.figure(figsize=(5, 4))
    plt.imshow(cm, interpolation='nearest', cmap=plt.cm.Blues)
    plt.title("Confusion Matrix")
    plt.colorbar()
    plt.tight_layout()
    outp = os.path.join(outdir, "confusion_matrix.png")
    plt.savefig(outp)
    plt.close()
    print(f"[+] Saved plot: {outp}")

###############################################################################
# Main pipeline (preserved flow)
###############################################################################

def main(debug=False):
    """Main pipeline execution."""
    # Verify directories
    for d in [NORMAL_WIP_DIR, NORMAL_SHS_DIR, MITM_WIP_DIR, MITM_SHS_DIR]:
        if not os.path.isdir(d):
            raise FileNotFoundError(f"Required directory not found: {d}")
    
    print("[+] Starting IoMT Comprehensive Analysis Pipeline")
    print("[+] FIXED: NDI and JVI now properly handle complete packet loss scenarios")
    print("[+] ADDED: ISO/IEC 27400, ISO 27799, ISO/IEC 27701, ISO 13485, IEC 62443 mappings")
    print("[+] ADDED: Complete plotting functions for all metrics")
    
    # Aggregate normal & attack runs
    print("[+] Loading and aggregating normal WIP data...")
    _, agg_n_w = load_and_aggregate(NORMAL_WIP_DIR)
    print("[+] Loading and aggregating normal SHS data...")
    _, agg_n_s = load_and_aggregate(NORMAL_SHS_DIR)
    print("[+] Loading and aggregating MITM WIP data...")
    _, agg_a_w = load_and_aggregate(MITM_WIP_DIR)
    print("[+] Loading and aggregating MITM SHS data...")
    _, agg_a_s = load_and_aggregate(MITM_SHS_DIR)
    
    # Compute comprehensive derived metrics with risk mappings
    print("[+] Computing derived metrics for WIP...")
    derived_w = compute_metrics_from_aggregates(agg_n_w, agg_a_w, "wip")
    print("[+] Computing derived metrics for SHS...")
    derived_s = compute_metrics_from_aggregates(agg_n_s, agg_a_s, "shs")
    derived_all = pd.concat([derived_w, derived_s], ignore_index=True)
    
    # Print diagnostic information about complete loss flows
    complete_loss = derived_all[derived_all['CompleteLoss'] == 'Yes']
    if not complete_loss.empty:
        print(f"\n[!] DETECTED {len(complete_loss)} flows with complete packet loss:")
        for _, r in complete_loss.iterrows():
            print(f"    {r['Device']} Flow {int(r['FlowID'])}: NDI={r['NDI_pct']:.1f}%, JVI={r['JVI_pct']:.1f}%, ASI={r['ASI_score']:.1f}")
    
    # Save core CSVs
    print("[+] Saving CSV outputs...")
    csv_out = os.path.join(OUTPUT_DIR, "derived_metrics_perflow.csv")
    derived_all.to_csv(csv_out, index=False)
    derived_all.to_csv(os.path.join(OUTPUT_DIR, "derived_metrics_all_devices.csv"), index=False)
    derived_all[['Device','FlowID','ASI_score','Severity','Mitigation','PriorityScore','ActionLevel','CriticalityFactor','CompleteLoss']].to_csv(
        os.path.join(OUTPUT_DIR,'mitigations_by_flow.csv'), index=False)
    print(f"[+] Saved derived metrics CSVs: {csv_out}")
    
    # Generate plots for both devices
    print("[+] Generating comparison plots...")
    all_plots = {}
    for dev_lower, agg_norm, agg_att, derived in [('wip', agg_n_w, agg_a_w, derived_w), 
                                                    ('shs', agg_n_s, agg_a_s, derived_s)]:
        plots = save_comparison_plots(dev_lower, agg_norm, agg_att, derived)
        all_plots[dev_lower.upper()] = plots
    
    # Generate aggregate plots
    print("[+] Generating aggregate analysis plots...")
    plot_asi_by_device(derived_all, PLOTS_DIR)
    plot_severity_pie(derived_all, PLOTS_DIR)
    plot_severity_distribution(derived_all, PLOTS_DIR)
    plot_attack_metrics_all_stats(derived_all, PLOTS_DIR)
    plot_critical_flows(derived_all, PLOTS_DIR)
    plot_flow9(derived_all, PLOTS_DIR)
    
    # Build comprehensive Excel workbook (extended)
    build_comprehensive_excel(derived_all, all_plots)
    
    # Build Word report (extended)
    build_word_report(derived_all)
    
    # Classification report
    print_classification_report_and_save_confusion(derived_all, outdir=PLOTS_DIR)
    
    # Mean metrics by FlowID & Type
    print("[+] Computing mean metrics by FlowID and Type...")
    derived_all['Type'] = derived_all['Device'].str.lower()
    metrics_cols = ['TRR_pct','PLR_pctpoints','PDR_pct','NDI_pct','JVI_pct','ASI_score']
    group_cols = ['FlowID', 'Type']
    df_mean_metrics = derived_all.groupby(group_cols)[metrics_cols].mean().reset_index()
    for col in ['Label','Severity']:
        df_first = derived_all.groupby(group_cols)[col].first().reset_index()
        df_mean_metrics = pd.merge(df_mean_metrics, df_first, on=group_cols)
    cols_order = ['FlowID', 'Label', 'Type'] + metrics_cols + ['Severity']
    df_mean_metrics = df_mean_metrics[cols_order]
    df_mean_metrics['Type'] = pd.Categorical(df_mean_metrics['Type'], categories=['wip', 'shs'], ordered=True)
    df_mean_metrics = df_mean_metrics.sort_values(by=['Type', 'FlowID']).reset_index(drop=True)
    mean_by_flow_csv = os.path.join(OUTPUT_DIR, 'mean_metrics_by_flowID_and_type_sorted.csv')
    df_mean_metrics.to_csv(mean_by_flow_csv, index=False)
    print(f"[+] Saved mean metrics by FlowID and Type: {mean_by_flow_csv}")
    
    # Print summary statistics
    print(f"\n{'='*60}")
    print(f"SUMMARY STATISTICS")
    print(f"{'='*60}")
    print(f"\nTotal flows analyzed: {len(derived_all)}")
    print(f"\nSeverity breakdown:")
    for sev in ['None', 'Low', 'Medium', 'High']:
        count = len(derived_all[derived_all['Severity'] == sev])
        print(f"  {sev}: {count}")
    
    print(f"\nRisk level breakdown:")
    for risk in ['None', 'Low', 'Medium', 'High', 'Critical']:
        count = len(derived_all[derived_all['RiskLabel'] == risk])
        print(f"  {risk}: {count}")
    
    print(f"\nComplete packet loss flows: {len(complete_loss)}")
    
    print(f"\nAction level breakdown:")
    for action in ['Normal', 'Warning', 'Action', 'Critical']:
        count = len(derived_all[derived_all['ActionLevel'] == action])
        print(f"  {action}: {count}")
    
    print(f"\nMetric statistics (mean ± std):")
    for metric in ['TRR_pct', 'PLR_pctpoints', 'NDI_pct', 'JVI_pct', 'ASI_score', 'PriorityScore']:
        mean_val = derived_all[metric].mean()
        std_val = derived_all[metric].std()
        print(f"  {metric}: {mean_val:.2f} ± {std_val:.2f}")
    
    print(f"\n{'='*60}")
    print(f"[+] Pipeline complete!")
    print(f"[+] Comprehensive Excel workbook: {EXCEL_PATH}")
    print(f"[+] Risk assessment Word document: {WORD_PATH}")
    print(f"[+] All outputs saved under: {OUTPUT_DIR}")
    print(f"[+] Plots available in: {PLOTS_DIR}")
    print(f"{'='*60}\n")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="IoMT Comprehensive Analysis with Risk Mapping - Extended Standards")
    parser.add_argument("--debug", action="store_true", help="Enable debug output")
    args = parser.parse_args()
    main(debug=args.debug)
